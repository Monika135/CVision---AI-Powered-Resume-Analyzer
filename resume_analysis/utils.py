import spacy
import io
import fitz
from docx import Document
nlp = spacy.load("en_core_web_md")


def extract_text_from_pdf(file_bytes):
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    return text


def extract_text_from_docx(file_path_or_bytes):
    if isinstance(file_path_or_bytes, bytes):
        doc = Document(io.BytesIO(file_path_or_bytes))
    else:
        doc = Document(file_path_or_bytes)
    full = []
    for txt in doc.paragraphs:
        full.append(txt.text)
    return "\n".join(full)


def clean_and_tokenize(text):
    doc = nlp(text)
    tokens = [token.lemma_.lower() for token in doc if token.is_alpha and not token.is_stop]
    return set(tokens)


def analyze_resume(resume_text, job_description):
    try:
        resume_doc = nlp(resume_text)
        jd_doc = nlp(job_description)
        score = resume_doc.similarity(jd_doc) * 100
    except Exception:
        score = 0.0

    resume_tokens = clean_and_tokenize(resume_text)
    jd_tokens = clean_and_tokenize(job_description)

    matched = sorted(list(resume_tokens & jd_tokens))
    missing = sorted(list(jd_tokens - resume_tokens))

    suggestions = [f"Consider adding details with '{m}'." for m in missing[:8]]

    return {
        "score": round(score, 2),
        "matched_keywords": matched[:20],
        "missing_keywords": missing[:20],
        "suggestions": suggestions,
    }